from docx import Document
from docx.shared import Cm

documento = Document()

# # Adicionar uma tabela
# tabela = documento.add_table(rows=3,cols=2)
# celula00 = tabela.cell(0,0)
# celula00.text = 'Nome'
# celula01 = tabela.cell(0,1)
# celula01.text = 'Idade'
# celula10 = tabela.cell(1,0)
# celula10.text = 'Rafael'
# celula11 = tabela.cell(1,1)
# celula11.text = '45'
# celula20 = tabela.cell(2,0)
# celula20.text = 'Amanda'
# celula21 = tabela.cell(2,1)
# celula21.text = '25'

# Preço de produtos
# dados a serem inseridos na tabela
registros = [
    [3,'101','Maça'],
    [7,'422','Ovos'],
    [4,'631','Banana']
]
# Criar a tabela
tabela = documento.add_table(rows=1,cols=3)
# Definindo o cabeçalho
cabecalho = tabela.rows[0].cells
cabecalho[0].text = 'Quantidade'
cabecalho[1].text = 'Id'
cabecalho[2].text = 'Descrição'

for quantidade, id, descricao in registros:
    linha_atual = tabela.add_row().cells
    linha_atual[0].text = str(quantidade)
    linha_atual[1].text = id
    linha_atual[2].text = descricao

documento.save('demo.docx')