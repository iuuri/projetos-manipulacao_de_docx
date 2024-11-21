from docx import Document
from docx.shared import Cm

documento = Document()
# para add um título
documento.add_heading('Título do documento', 0)
# para add um parágrafo
paragrafo = documento.add_paragraph('Um parágrafo simples')
paragrafo.add_run(' e super importante ').bold = True
paragrafo.add_run('do autor ')
paragrafo.add_run('jhonatan').italic = True

# adicionar heading(cabeçalho)
documento.add_heading('Título nível 1', level=1)
documento.add_heading('Título nível 2', level=2)
documento.add_heading('Título nível 3', level=3)
documento.add_heading('Título nível 4', level=4)

# Formatação de estilo
documento.add_paragraph('Formatação "No Spacing"', style='No Spacing')
documento.add_paragraph('Formatação "Heading1"', style='Heading 1')
documento.add_paragraph('Formatação "Heading 2"', style='Heading 2')
documento.add_paragraph('Formatação "Heading 3"', style='Heading 3')
documento.add_page_break()
documento.add_paragraph('Formatação "Title"', style='Title')
documento.add_paragraph('Formatação "Subtitle"', style='Subtitle')
documento.add_paragraph('Formatação "Quote"', style='Quote')
documento.add_paragraph('Formatação "Intense Quote"', style='Intense Quote')
documento.add_paragraph('Formatação "List Paragraph"', style='List Paragraph')
documento.add_paragraph(
    'Primeiro item em uma lista com pontos', style='List Bullet')
documento.add_paragraph(
    'primeiro item em uma lista numerada', style='List Number')
# Adicionar uma imagem
documento.add_picture('notebook.jpg', width=Cm(5.25))
#Quebra de linha
documento.add_page_break()
# Adicionar uma tabela
# tabela = documento.add_table(rows=3,cols=2)
# celula00 = tabela.cell(0,0)
# celula00.text = 'Nome'
# celula01 = tabela.cell(0,1)
# celula01.text = 'Idade'
# celula10 = tabela.cell(1,0)
# celula10.text = 'Rafael'
# celula11 = tabela.cell(1,1)
# celula11.text = '45'
# celula20 = tabela.cell(2,0)
# celula20.text = 'Amanda'
# celula21 = tabela.cell(2,1)
# celula21.text = '25'

# Preço de produtos
# dados a serem inseridos na tabela
registros = [
    [3, '101', 'Maça'],
    [7, '422', 'Ovos'],
    [4, '631', 'Banana']
]
# Criar a tabela
tabela = documento.add_table(rows=1, cols=3)
# Definindo o cabeçalho
cabecalho = tabela.rows[0].cells
cabecalho[0].text = 'Quantidade'
cabecalho[1].text = 'Id'
cabecalho[2].text = 'Descrição'

for quantidade, id, descricao in registros:
    linha_atual = tabela.add_row().cells
    linha_atual[0].text = str(quantidade)
    linha_atual[1].text = id
    linha_atual[2].text = descricao

documento.save('demo.docx')